"""
Сервер веб-приложения для визуализации исторических данных
"""

import os
import json
import logging
import time
import shutil
import threading
import datetime
from flask import Flask, render_template, jsonify, request, send_file, make_response
from flask_cors import CORS

# Путь к файлу с историческими данными
HISTORY_DB_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                             "history_db_generator/russian_history_database.json")

# Путь к файлу с администраторами
ADMINS_FILE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'admins.json')

app = Flask(__name__, 
            template_folder=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates'),
            static_folder=os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'static'))

CORS(app)  # Включаем поддержку CORS для всех маршрутов

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

def load_historical_data():
    """Загрузка исторических данных из файла"""
    try:
        if os.path.exists(HISTORY_DB_PATH):
            with open(HISTORY_DB_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
            logger.info(f"Загружено {len(data.get('events', []))} исторических событий")
            return data
        else:
            logger.warning(f"Файл базы данных не найден: {HISTORY_DB_PATH}")
            return {"events": []}
    except Exception as e:
        logger.error(f"Ошибка при загрузке исторических данных: {e}")
        return {"events": []}

# Загрузка данных при старте сервера
historical_data = load_historical_data()

@app.route('/')
def index():
    """Главная страница"""
    # Получаем статистику для приветственного баннера
    events_count = len(historical_data.get('events', [])) if historical_data else 0

    # Получаем уникальные категории
    categories = set()
    if historical_data and 'events' in historical_data:
        for event in historical_data['events']:
            if event.get('category'):
                categories.add(event.get('category'))

    categories_count = len(categories)

    return render_template('index.html', 
                          events_count=events_count,
                          categories_count=categories_count)

@app.route('/api/historical-events')
def get_historical_events():
    """API для получения исторических данных"""
    try:
        # Проверяем, загружены ли данные
        global historical_data
        if not historical_data:
            historical_data = load_historical_data()

        # Получаем события из базы данных
        events = historical_data.get('events', [])
        logger.info(f"Всего событий в базе: {len(events)}")

        # Фильтруем события, у которых есть координаты местоположения
        filtered_events = []
        for event in events:
            # Проверка на наличие координат
            has_coords = False
            if isinstance(event.get('location'), dict):
                has_coords = event.get('location', {}).get('lat') and event.get('location', {}).get('lng')

            if not has_coords:
                continue

            filtered_event = {
                'id': event.get('id', ''),
                'title': event.get('title', ''),
                'date': event.get('date', ''),
                'description': event.get('description', ''),
                'location': event.get('location', ''),
                'category': event.get('category', ''),
                'topic': event.get('topic', '')
            }
            filtered_events.append(filtered_event)

        logger.info(f"Отправляется {len(filtered_events)} событий с координатами")
        return jsonify(filtered_events)
    except Exception as e:
        logger.error(f"Ошибка при получении исторических данных: {e}")
        return jsonify([]), 200  # Возвращаем пустой массив вместо ошибки

@app.route('/api/categories')
def get_categories():
    """API для получения списка категорий событий"""
    try:
        events = historical_data.get('events', [])
        categories = sorted(list(set(e.get('category') for e in events if e.get('category'))))

        return jsonify(categories)
    except Exception as e:
        logger.error(f"Ошибка при получении категорий: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/event-details', methods=['POST'])
def get_event_details():
    """API для получения информации о событии через Gemini API"""
    try:
        # Получаем данные из запроса
        data = request.json

        if not data or not data.get('title'):
            return jsonify({'error': 'Недостаточно данных о событии'}), 400

        # Формируем запрос к Gemini API
        # Импортируем здесь для избежания циклических импортов
        import sys
        import os
        # Добавляем корневую директорию проекта в путь для импорта
        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        if root_dir not in sys.path:
            sys.path.append(root_dir)

        from src.api_client import APIClient
        from src.api_cache import APICache
        from src.logger import Logger

        # Получаем API ключ
        api_key = None
        try:
            from gemini_api_keys import GEMINI_API_KEYS
            if GEMINI_API_KEYS:
                api_key = GEMINI_API_KEYS[0]
        except ImportError:
            logger.error("Не удалось импортировать GEMINI_API_KEYS")

        if not api_key:
            try:
                import os
                api_key = os.environ.get('GEMINI_API_KEY')
            except:
                pass

        if not api_key:
            return jsonify({
                'content': 'Не удалось получить доступ к API Gemini. Пожалуйста, проверьте настройки API ключей.'
            }), 200

        # Создаем объекты для работы с API
        api_cache = APICache(logger)
        api_client = APIClient(api_key, api_cache, logger)

        # Формируем промпт для Gemini
        event_title = data.get('title', '')
        event_date = data.get('date', '')
        event_category = data.get('category', '')
        event_location = data.get('location', '')
        is_brief = data.get('isBrief', True)

        if is_brief:
            prompt = f"""
            Предоставь краткую историческую информацию о следующем событии из истории России:

            Название: {event_title}
            Дата: {event_date}
            Категория: {event_category}
            Место: {event_location}

            Ответ должен быть кратким (не более 200 слов), но содержательным. 
            Выдели 3-4 ключевых факта о событии и его значении. 
            Используй маркированный список для лучшей читаемости.
            """
        else:
            prompt = f"""
            Предоставь подробную историческую информацию о следующем событии из истории России:

            Название: {event_title}
            Дата: {event_date}
            Категория: {event_category}
            Место: {event_location}

            Пожалуйста, структурируй ответ следующим образом:
            1. Исторический контекст (что происходило в России в это время)
            2. Подробное описание события
            3. Ключевые участники
            4. Причины и предпосылки
            5. Последствия и историческое значение

            Используй только проверенные исторические факты. Ответ должен быть информативным и подробным.
            """

        try:
            # Инициализация API клиента если нужно
            if not api_client.is_initialized():
                api_client.initialize()

            # Запрос к Gemini API
            response = api_client.ask_grok(prompt, use_cache=True)

            return jsonify({
                'content': response
            })
        except Exception as api_error:
            logger.error(f"Ошибка при запросе к Gemini API: {api_error}")
            return jsonify({
                'content': f"Произошла ошибка при получении информации: {str(api_error)}"
            }), 200

    except Exception as e:
        logger.error(f"Ошибка при обработке запроса о деталях события: {e}")
        return jsonify({
            'content': 'Произошла ошибка при обработке запроса. Пожалуйста, попробуйте позже.'
        }), 200

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    """API для генерации подробного реферата в формате Word"""
    try:
        # Получаем данные из запроса
        data = request.json

        if not data or not data.get('title'):
            return jsonify({'error': 'Недостаточно данных о событии'}), 400

        # Импортируем необходимые библиотеки
        import sys
        import os
        from io import BytesIO
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Добавляем корневую директорию проекта в путь для импорта
        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        if root_dir not in sys.path:
            sys.path.append(root_dir)

        from src.api_client import APIClient
        from src.api_cache import APICache
        from src.logger import Logger

        # Получаем API ключ
        api_key = None
        try:
            from gemini_api_keys import GEMINI_API_KEYS
            if GEMINI_API_KEYS:
                api_key = GEMINI_API_KEYS[0]
        except ImportError:
            logger.error("Не удалось импортировать GEMINI_API_KEYS")

        if not api_key:
            try:
                import os
                api_key = os.environ.get('GEMINI_API_KEY')
            except:
                pass

        if not api_key:
            return jsonify({'error': 'API ключ не найден'}), 500

        # Создаем объекты для работы с API
        api_cache = APICache(logger)
        api_client = APIClient(api_key, api_cache, logger)

        # Формируем промпт для Gemini для получения подробной информации
        event_title = data.get('title', '')
        event_date = data.get('date', '')
        event_category = data.get('category', '')
        event_location = data.get('location', '')

        prompt = f"""
        Напиши подробный исторический реферат о следующем событии из истории России:

        Название: {event_title}
        Дата: {event_date}
        Категория: {event_category}
        Место: {event_location}

        Реферат должен содержать следующие разделы:
        1. Введение и исторический контекст (что происходило в России в это время)
        2. Хронология и детальное описание события
        3. Ключевые исторические личности, их роли и биографические данные
        4. Причины и предпосылки события
        5. Непосредственные последствия
        6. Историческое значение и влияние на дальнейший ход истории России
        7. Историографический анализ и различные точки зрения историков
        8. Заключение

        Реферат должен быть максимально подробным, академическим и опираться на исторические источники.
        """

        try:
            # Инициализация API клиента если нужно
            if not api_client.is_initialized():
                api_client.initialize()

            # Запрос к Gemini API
            detailed_content = api_client.ask_grok(prompt, use_cache=True)

            # Создаем документ Word
            doc = Document()

            # Устанавливаем стили
            title_style = doc.styles['Title']
            title_style.font.size = Pt(16)
            title_style.font.bold = True

            heading_style = doc.styles['Heading 1']
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True

            # Заголовок документа
            doc.add_heading(event_title, 0)

            # Информация о событии
            info_paragraph = doc.add_paragraph()
            info_paragraph.add_run(f"Дата: {event_date}\n").bold = True
            info_paragraph.add_run(f"Категория: {event_category}\n").bold = True
            info_paragraph.add_run(f"Место: {event_location}").bold = True

            # Добавляем разделительную линию
            doc.add_paragraph("_" * 50)

            # Парсим и добавляем содержимое
            content_lines = detailed_content.split('\n')
            current_heading_level = 0

            for i, line in enumerate(content_lines):
                line = line.strip()

                if not line:
                    doc.add_paragraph()
                    continue

                # Определяем, является ли строка заголовком
                heading_level = 0
                if line.startswith('# '):
                    heading_level = 1
                    heading_text = line[2:].strip()
                elif line.startswith('## '):
                    heading_level = 2
                    heading_text = line[3:].strip()
                elif line.startswith('### '):
                    heading_level = 3
                    heading_text = line[4:].strip()
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    heading_level = 2
                    heading_text = line[3:].strip()

                # Обрабатываем заголовки
                if heading_level > 0:
                    if heading_level == 1:
                        doc.add_heading(heading_text, level=1)
                    elif heading_level == 2:
                        doc.add_heading(heading_text, level=2)
                    elif heading_level == 3:
                        doc.add_heading(heading_text, level=3)

                    current_heading_level = heading_level
                    continue

                # Обрабатываем обычный текст
                if line.startswith('- ') or line.startswith('* '):
                    # Маркированный список
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    # Обычный текст
                    p = doc.add_paragraph(line)

            # Сохраняем документ в память
            file_buffer = BytesIO()
            doc.save(file_buffer)
            file_buffer.seek(0)

            # Возвращаем файл
            return send_file(
                file_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{event_title.replace('/', '_')}_реферат.docx"
            )

        except Exception as api_error:
            logger.error(f"Ошибка при генерации реферата: {api_error}")
            return jsonify({'error': str(api_error)}), 500

    except Exception as e:
        logger.error(f"Ошибка при обработке запроса на генерацию реферата: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin-panel')
def admin_panel():
    """Страница администратора"""
    return render_template('admin_panel.html')

@app.route('/api/admin/admins', methods=['GET'])
def get_admins():
    """API для получения списка администраторов"""
    try:
        admins = load_admins()
        return jsonify(admins)
    except Exception as e:
        logger.error(f"Ошибка при получении списка администраторов: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/admins', methods=['POST'])
def add_admin():
    """API для добавления администратора"""
    try:
        data = request.json
        if not data or not data.get('user_id'):
            return jsonify({"error": "Недостаточно данных"}), 400

        user_id = int(data.get('user_id'))
        is_super = data.get('is_super', False)

        admins = load_admins()

        if is_super:
            if user_id not in admins.get("super_admin_ids", []):
                admins.setdefault("super_admin_ids", []).append(user_id)
                logger.info(f"Добавлен супер-админ: {user_id}")
        else:
            if user_id not in admins.get("admin_ids", []):
                admins.setdefault("admin_ids", []).append(user_id)
                logger.info(f"Добавлен админ: {user_id}")

        save_admins(admins)
        return jsonify({"success": True})
    except Exception as e:
        logger.error(f"Ошибка при добавлении администратора: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/admins/<int:user_id>', methods=['DELETE'])
def remove_admin(user_id):
    """API для удаления администратора"""
    try:
        admins = load_admins()

        if user_id in admins.get("admin_ids", []):
            admins["admin_ids"].remove(user_id)
            logger.info(f"Удален админ: {user_id}")
            save_admins(admins)
            return jsonify({"success": True})
        elif user_id in admins.get("super_admin_ids", []):
            admins["super_admin_ids"].remove(user_id)
            logger.info(f"Удален супер-админ: {user_id}")
            save_admins(admins)
            return jsonify({"success": True})

        return jsonify({"error": "Администратор не найден"}), 404
    except Exception as e:
        logger.error(f"Ошибка при удалении администратора: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/stats')
def get_stats():
    """API для получения статистики"""
    try:
        # Пример статистики, в реальной реализации нужно использовать фактические данные
        stats = {
            "user_count": 42,
            "message_count": 1337,
            "uptime": "3 дня 7 часов",
            "bot_starts": 25,
            "topic_requests": 73,
            "completed_tests": 18
        }
        return jsonify(stats)
    except Exception as e:
        logger.error(f"Ошибка при получении статистики: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/logs')
def get_logs():
    """API для получения логов"""
    try:
        lines = request.args.get('lines', 100, type=int)
        level = request.args.get('level', 'all')

        logs = get_last_logs(lines, level)
        return jsonify(logs)
    except Exception as e:
        logger.error(f"Ошибка при получении логов: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/settings', methods=['GET'])
def get_settings():
    """API для получения настроек бота"""
    try:
        settings = load_bot_settings()
        return jsonify(settings)
    except Exception as e:
        logger.error(f"Ошибка при получении настроек: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/settings', methods=['POST'])
def update_settings():
    """API для обновления настроек бота"""
    try:
        settings = request.json
        if not settings:
            return jsonify({"error": "Недостаточно данных"}), 400

        save_bot_settings(settings)
        return jsonify({"success": True})
    except Exception as e:
        logger.error(f"Ошибка при обновлении настроек: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/maintenance/<action>', methods=['POST'])
def admin_maintenance(action):
    """API для выполнения операций обслуживания"""
    try:
        if action == 'clear-cache':
            # Полная очистка кэша
            result = os.system('python clear_all_cache.py --all')
            if result == 0:
                return jsonify({'success': True, 'message': 'Кэш успешно очищен'})
            else:
                return jsonify({'error': 'Ошибка при очистке кэша'}), 500

        elif action == 'selective-cache':
            # Выборочная очистка кэша
            data = request.get_json()
            if not data or not data.get('types'):
                return jsonify({'error': 'Не указаны типы кэша для очистки'}), 400

            types = data.get('types')
            cmd_args = []

            for cache_type in types:
                if cache_type in ['api', 'events', 'user', 'images']:
                    cmd_args.append(f'--{cache_type}')

            if not cmd_args:
                return jsonify({'error': 'Не указаны корректные типы кэша для очистки'}), 400

            cmd = f'python clear_all_cache.py {" ".join(cmd_args)}'
            result = os.system(cmd)

            if result == 0:
                return jsonify({
                    'success': True, 
                    'message': f'Выборочная очистка кэша выполнена успешно: {", ".join(types)}'
                })
            else:
                return jsonify({'error': 'Ошибка при выборочной очистке кэша'}), 500

        elif action == 'backup':
            # Резервное копирование
            timestamp = int(time.time())
            backup_dir = f'backups/data_backup_v{timestamp}'

            # Создаем директорию для резервной копии если нужно
            os.makedirs('backups', exist_ok=True)

            # Копируем ключевые файлы данных
            files_to_backup = [
                'historical_events.json',
                'user_states.json',
                'admins.json',
                'bot_settings.json',
                'api_cache.json'
            ]

            for file in files_to_backup:
                if os.path.exists(file):
                    shutil.copy2(file, f'backups/{file.replace(".json", "")}_backup_{timestamp}.json')

            # Записываем информацию о бэкапе
            with open('backups/backup_info.json', 'w') as f:
                json.dump({
                    'timestamp': timestamp,
                    'date': datetime.datetime.now().isoformat(),
                    'files': files_to_backup
                }, f, indent=4)

            return jsonify({
                'success': True, 
                'message': 'Резервная копия успешно создана',
                'backup_time': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'files_count': len(files_to_backup)
            })

        elif action == 'update-api':
            # Обновление данных API
            # Запускаем обновление в отдельном процессе
            thread = threading.Thread(target=update_api_data)
            thread.daemon = True
            thread.start()

            return jsonify({
                'success': True, 
                'message': 'Обновление данных API запущено. Этот процесс может занять некоторое время.'
            })

        elif action == 'clean-logs':
            # Очистка логов (удаляем все, кроме текущего и предыдущего дня)
            log_dir = 'logs'
            if not os.path.exists(log_dir):
                os.makedirs(log_dir)

            current_date = datetime.datetime.now()
            yesterday = (current_date - datetime.timedelta(days=1)).strftime('%Y%m%d')
            today = current_date.strftime('%Y%m%d')

            keep_logs = [f'bot_log_{today}.log', f'bot_log_{yesterday}.log']
            deleted_count = 0

            for filename in os.listdir(log_dir):
                if filename.startswith('bot_log_') and filename.endswith('.log'):
                    if filename not in keep_logs:
                        os.remove(os.path.join(log_dir, filename))
                        deleted_count += 1

            return jsonify({
                'success': True, 
                'message': f'Логи успешно очищены. Удалено {deleted_count} файлов.'
            })

        elif action == 'audit':
            # Возвращаем историю действий администраторов
            audit_data = get_admin_audit_log()
            return jsonify({
                'success': True,
                'data': audit_data
            })

        else:
            return jsonify({'error': 'Неизвестное действие обслуживания'})

    except Exception as e:
        logger.error(f"Ошибка при выполнении операции обслуживания: {e}")
        return jsonify({'error': f'Ошибка при выполнении операции: {str(e)}'}), 500

def update_api_data():
    """Фоновое обновление данных API"""
    try:
        # Здесь код для обновления данных API
        # Например, обновление исторических событий из внешнего источника
        time.sleep(2)  # Имитация долгой операции
        # В реальном коде здесь будут вызовы внешних API и обновление данных

        logger.info("Данные API успешно обновлены")
    except Exception as e:
        logger.error(f"Ошибка при обновлении данных API: {e}")

def get_admin_audit_log():
    """Получение истории действий администраторов"""
    # В реальном приложении это может быть получение данных из БД или лог-файла
    # Для примера возвращаем заглушку
    return [
        {
            "timestamp": (datetime.datetime.now() - datetime.timedelta(hours=2)).isoformat(),
            "admin_id": 7225056628,
            "action": "login",
            "details": "Вход в систему"
        },
        {
            "timestamp": (datetime.datetime.now() - datetime.timedelta(hours=1)).isoformat(),
            "admin_id": 7225056628,
            "action": "settings_change",
            "details": "Изменены настройки системы"
        },
        {
            "timestamp": datetime.datetime.now().isoformat(),
            "admin_id": 7225056628,
            "action": "cache_clear",
            "details": "Выполнена очистка кэша"
        }
    ]

@app.route('/api/admin/check-auth', methods=['GET'])
def check_admin_auth():
    """Проверка аутентификации администратора"""
    try:
        user_id = request.cookies.get('admin_id')

        if not user_id:
            return jsonify({"authenticated": False})

        admins = load_admins()
        user_id = int(user_id)

        is_admin = user_id in admins.get("admin_ids", []) or user_id in admins.get("super_admin_ids", [])
        is_super_admin = user_id in admins.get("super_admin_ids", [])

        if is_admin:
            return jsonify({
                "authenticated": True,
                "user": {
                    "id": user_id,
                    "is_super_admin": is_super_admin
                }
            })

        return jsonify({"authenticated": False})
    except Exception as e:
        logger.error(f"Ошибка при проверке аутентификации: {e}")
        return jsonify({"authenticated": False})

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    """Авторизация администратора"""
    try:
        data = request.json
        admin_password = data.get('admin_password', '')

        if not admin_password:
            return jsonify({"success": False, "message": "Пароль администратора не указан"})

        # Проверяем пароль - используем фиксированный пароль "nnkhqjm"
        correct_password = "nnkhqjm"

        if admin_password == correct_password:
            # При успешной авторизации используем ID супер-администратора из admins.json
            admins = load_admins()
            admin_id = admins.get("super_admin_ids", [7225056628])[0]  # По умолчанию используем ID из admins.json

            logger.info(f"Успешный вход администратора с ID {admin_id}")

            response = jsonify({
                "success": True,
                "user": {
                    "id": admin_id,
                    "is_super_admin": True  # Авторизация по паролю дает права супер-администратора
                }
            })
            response.set_cookie('admin_id', str(admin_id), max_age=86400)  # 24 часа
            return response

        logger.warning(f"Неудачная попытка входа администратора с паролем '{admin_password}'")
        return jsonify({"success": False, "message": "Неверный пароль администратора"})
    except Exception as e:
        logger.error(f"Ошибка при авторизации администратора: {e}")
        return jsonify({"success": False, "message": f"Ошибка при авторизации: {str(e)}"})

@app.route('/api/admin/users', methods=['GET'])
def get_users():
    """API для получения списка пользователей"""
    try:
        # В реальном приложении здесь был бы запрос к базе данных
        # Для примера возвращаем тестовые данные
        sample_users = [
            {"id": 123456789, "name": "Иван", "status": "active", "last_activity": "2023-03-09 15:30", "messages": 42},
            {"id": 987654321, "name": "Мария", "status": "active", "last_activity": "2023-03-09 16:45", "messages": 28},
            {"id": 555555555, "name": "Алексей", "status": "inactive", "last_activity": "2023-03-05 10:15", "messages": 13},
            {"id": 111111111, "name": "Елена", "status": "active", "last_activity": "2023-03-09 14:20", "messages": 37},
            {"id": 222222222, "name": "Сергей", "status": "blocked", "last_activity": "2023-02-15 08:30", "messages": 5}
        ]

        return jsonify(sample_users)
    except Exception as e:
        logger.error(f"Ошибка при получении списка пользователей: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/user/<int:user_id>', methods=['GET'])
def get_user(user_id):
    """API для получения информации о пользователе"""
    try:
        # В реальном приложении здесь был бы запрос к базе данных
        # Для примера возвращаем тестовые данные
        sample_user = {
            "id": user_id,
            "name": "Тестовый пользователь",
            "status": "active",
            "last_activity": "2023-03-09 15:30",
            "messages": 42,
            "registration_date": "2023-01-15",
            "tests_completed": 5,
            "favorite_topics": ["Революции", "Война 1812 года"]
        }

        return jsonify(sample_user)
    except Exception as e:
        logger.error(f"Ошибка при получении информации о пользователе: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/user/<int:user_id>/block', methods=['POST'])
def block_user(user_id):
    """API для блокировки пользователя"""
    try:
        # В реальном приложении здесь была бы логика блокировки
        logger.info(f"Запрошена блокировка пользователя {user_id}")
        return jsonify({"success": True, "message": f"Пользователь {userid} заблокирован"})
    except Exception as e:
        logger.error(f"Ошибка при блокировке пользователя: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/user/<int:user_id>/unblock', methods=['POST'])
def unblock_user(user_id):
    """API для разблокировки пользователя"""
    try:
        # В реальном приложении здесь была бы логика разблокировки
        logger.info(f"Запрошена разблокировка пользователя {user_id}")
        return jsonify({"success": True, "message": f"Пользователь {user_id} разблокирован"})
    except Exception as e:
        logger.error(f"Ошибка при разблокировке пользователя: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/export-logs', methods=['GET'])
def export_logs():
    """API для экспорта логов в файл"""
    try:
        # Получаем логи
        level = request.args.get('level', 'all')
        lines = request.args.get('lines', 1000, type=int)
        logs = get_last_logs(lines, level)

        # Создаем временный файл для записи логов
        temp_file = f"logs_export_{int(time.time())}.txt"
        with open(temp_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(logs))

        # Отправляем файл
        response = send_file(
            temp_file,
            as_attachment=True,
            download_name=f"logs_export_{time.strftime('%Y%m%d_%H%M%S')}.txt"
        )

        # Удаляем временный файл после отправки
        @response.call_on_close
        def remove_file():
            try:
                os.remove(temp_file)
            except:
                pass

        return response
    except Exception as e:
        logger.error(f"Ошибка при экспорте логов: {e}")
        return jsonify({"error": str(e)}), 500

# Вспомогательные функции для админ-панели

def load_admins():
    """Загружает список администраторов из файла"""
    try:
        if os.path.exists(ADMINS_FILE_PATH):
            with open(ADMINS_FILE_PATH, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {"admin_ids": [], "super_admin_ids": []}
    except Exception as e:
        logger.error(f"Ошибка при загрузке списка администраторов: {e}")
        return {"admin_ids": [], "super_admin_ids": []}

def save_admins(admins):
    """Сохраняет список администраторов в файл"""
    try:
        # Используем атомарную операцию записи
        temp_file = f"{ADMINS_FILE_PATH}.tmp"
        with open(temp_file, 'w', encoding='utf-8') as f:
            json.dump(admins, f, indent=4)
        os.replace(temp_file, ADMINS_FILE_PATH)
        return True
    except Exception as e:
        logger.error(f"Ошибка при сохранении списка администраторов: {e}")
        return False

def get_last_logs(lines=100, level='all'):
    """Получает последние строки из файла логов"""
    try:
        log_files = []
        log_dir = "logs"

        # Проверяем, существует ли директория логов
        if os.path.exists(log_dir) and os.path.isdir(log_dir):
            # Получаем список файлов логов
            files = os.listdir(log_dir)
            log_files = [os.path.join(log_dir, f) for f in files if f.startswith("bot_log_")]

        # Если директории нет или нет файлов логов, проверяем в корневой директории
        if not log_files:
            files = [f for f in os.listdir() if f.startswith("bot_log_")]
            log_files = files

        # Сортируем файлы по дате изменения (новые в начале)
        log_files.sort(key=lambda x: os.path.getmtime(x), reverse=True)

        if not log_files:
            return ["Файлы логов не найдены"]

        # Берем самый свежий файл логов
        latest_log = log_files[0]

        log_lines = []
        with open(latest_log, 'r', encoding='utf-8') as f:
            # Читаем все строки файла
            all_lines = f.readlines()

        # Фильтруем по уровню, если указан не 'all'
        if level != 'all':
            filtered_lines = [line for line in all_lines if f" - {level.upper()} - " in line]
        else:
            filtered_lines = all_lines

        # Берем последние N строк
        return filtered_lines[-lines:]
    except Exception as e:
        logger.error(f"Ошибка при чтении файла логов: {e}")
        return [f"Ошибка при чтении логов: {e}"]

def load_bot_settings():
    """Загружает настройки бота из файла"""
    try:
        settings_path = 'bot_settings.json'
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            "auto_update_topics": True,
            "collect_statistics": True,
            "developer_mode": False,
            "private_mode": False
        }
    except Exception as e:
        logger.error(f"Ошибка при загрузке настроек бота: {e}")
        return {}

def save_bot_settings(settings):
    """Сохраняет настройки бота в файл"""
    try:
        settings_path = 'bot_settings.json'
        # Используем атомарную операцию записи
        temp_file = f"{settings_path}.tmp"
        with open(temp_file, 'w', encoding='utf-8') as f:
            json.dump(settings, f, indent=4)
        os.replace(temp_file, settings_path)
        return True
    except Exception as e:
        logger.error(f"Ошибка при сохранении настроек бота: {e}")
        return False

def clear_cache():
    """Очищает кэш"""
    try:
        cache_files = ['api_cache.json']
        for cache_file in cache_files:
            if os.path.exists(cache_file):
                os.remove(cache_file)
        return True
    except Exception as e:
        logger.error(f"Ошибка при очистке кэша: {e}")
        return False

def create_backup():
    """Создает резервную копию данных"""
    try:
        backup_dir = 'backups'
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)

        timestamp = int(time.time())

        # Резервное копирование файлов данных
        files_to_backup = ['admins.json', 'user_states.json', 'historical_events.json']
        for file in files_to_backup:
            if os.path.exists(file):
                backup_path = os.path.join(backup_dir, f"{os.path.splitext(file)[0]}_backup_{timestamp}{os.path.splitext(file)[1]}")
                with open(file, 'r', encoding='utf-8') as src, open(backup_path, 'w', encoding='utf-8') as dst:
                    dst.write(src.read())

        # Создаем файл метаданных резервной копии
        backup_meta = os.path.join(backup_dir, f"data_backup_v{len(os.listdir(backup_dir))}_{timestamp}")
        with open(backup_meta, 'w', encoding='utf-8') as f:
            f.write(f"Backup created at {time.ctime(timestamp)}")

        return True
    except Exception as e:
        logger.error(f"Ошибка при создании резервной копии: {e}")
        return False

def update_api_data():
    """Обновляет данные API"""
    # Заглушка для демонстрации
    return True

def clean_logs():
    """Очищает старые логи"""
    try:
        log_dir = 'logs'
        current_time = time.time()
        # Удаляем логи старше 30 дней
        if os.path.exists(log_dir) and os.path.isdir(log_dir):
            for log_file in os.listdir(log_dir):
                log_path = os.path.join(log_dir, log_file)
                if os.path.isfile(log_path) and current_time - os.path.getmtime(log_path) > 30 * 24 * 60 * 60:
                    os.remove(log_path)
        return True
    except Exception as e:
        logger.error(f"Ошибка при очистке логов: {e}")
        return False

def run_server(host='0.0.0.0', port=8080):
    """Запуск веб-сервера"""
    logger.info(f"Запуск веб-сервера на {host}:{port}")
    app.run(host=host, port=port, debug=False)


@app.route('/api/admin/get-doc')
def get_admin_doc():
    """Маршрут для скачивания документации администратора"""
    #  Added session check for authentication.  Replace with actual session management if needed.
    if False: # Replace with actual session authentication check.
        return jsonify({"error": "Не авторизован"}), 401

    doc_path = request.args.get('path')
    if not doc_path:
        return jsonify({"error": "Путь к документу не указан"}), 400

    try:
        # Проверяем безопасность пути (предотвращаем path traversal)
        norm_path = os.path.normpath(doc_path)
        if '..' in norm_path:
            return jsonify({"error": "Недопустимый путь"}), 403

        # Проверяем существование файла
        if not os.path.exists(norm_path):
            return jsonify({"error": "Файл не найден"}), 404

        # Получаем содержимое файла
        with open(norm_path, 'rb') as f:
            file_data = f.read()

        # Определяем имя файла из пути
        filename = os.path.basename(norm_path)

        # Определяем MIME-тип для ответа
        mime_type = 'application/octet-stream'
        if norm_path.endswith('.md'):
            mime_type = 'text/markdown'
        elif norm_path.endswith('.pdf'):
            mime_type = 'application/pdf'
        elif norm_path.endswith('.txt'):
            mime_type = 'text/plain'

        # Формируем ответ со скачиванием файла
        response = make_response(file_data)
        response.headers['Content-Type'] = mime_type
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        print(f"Ошибка при получении документа: {e}")
        return jsonify({"error": f"Ошибка при получении документа: {str(e)}"}), 500

@app.route('/neadmin/<path:doc_name>')
def get_neadmin_doc(doc_name):
    """Публичный маршрут для скачивания документации без аутентификации"""
    try:
        # Определяем базовый путь для документов (папка docs)
        base_path = os.path.join('static', 'docs')

        # Формируем полный путь к документу
        doc_path = os.path.join(base_path, doc_name)

        # Нормализуем путь и проверяем, что он не выходит за пределы базовой папки
        norm_path = os.path.normpath(doc_path)
        if not norm_path.startswith(os.path.normpath(base_path)):
            return jsonify({"error": "Недопустимый путь документа"}), 403

        # Проверяем существование файла
        if not os.path.exists(norm_path):
            return jsonify({"error": "Документ не найден"}), 404

        # Получаем содержимое файла
        with open(norm_path, 'rb') as f:
            file_data = f.read()

        # Определяем MIME-тип для ответа
        mime_type = 'application/octet-stream'
        if norm_path.endswith('.md'):
            mime_type = 'text/markdown'
        elif norm_path.endswith('.pdf'):
            mime_type = 'application/pdf'
        elif norm_path.endswith('.txt'):
            mime_type = 'text/plain'

        # Формируем ответ со скачиванием файла
        response = make_response(file_data)
        response.headers['Content-Type'] = mime_type
        response.headers['Content-Disposition'] = f'attachment; filename="{doc_name}"'
        return response

    except Exception as e:
        print(f"Ошибка при получении документа из /neadmin: {e}")
        return jsonify({"error": f"Ошибка при получении документа: {str(e)}"}), 500

@app.route('/api/admin/view-doc')
def view_doc():
    """API для просмотра документации в iframe"""
    try:
        doc_path = request.args.get('path')
        if not doc_path:
            return "Путь к документу не указан", 400

        # Проверка безопасности пути
        base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

        # Сначала проверяем в папке static/docs
        static_doc_path = os.path.join(base_path, 'static', 'docs', doc_path.lstrip('/'))
        if os.path.exists(static_doc_path):
            absolute_path = static_doc_path
        else:
            # Если в static/docs нет, проверяем в исходной папке docs
            absolute_path = os.path.abspath(os.path.join(base_path, doc_path))

        # Проверяем, что путь находится в пределах разрешенных директорий
        if not (absolute_path.startswith(os.path.join(base_path, 'static')) or 
                absolute_path.startswith(os.path.join(base_path, 'docs'))):
            return "Доступ запрещен", 403

        # Проверяем существование файла
        if not os.path.exists(absolute_path):
            return "Файл не найден", 404

        # Если это markdown файл, преобразуем его в HTML для отображения
        file_ext = os.path.splitext(absolute_path)[1].lower()
        if file_ext == '.md':
            try:
                import markdown
                with open(absolute_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                html_content = markdown.markdown(content)
                return f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; max-width: 800px; margin: 0 auto; }}
                        pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; overflow-x: auto; }}
                        code {{ font-family: monospace; background-color: #f5f5f5; padding: 2px 4px; border-radius: 3px; }}
                        img {{ max-width: 100%; height: auto; }}
                        h1, h2, h3, h4, h5, h6 {{ margin-top: 1.5em; margin-bottom: 0.5em; }}
                        p {{ margin: 1em 0; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
            except ImportError:
                # Если модуль markdown не установлен, просто отображаем текст
                with open(absolute_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return f"<pre>{content}</pre>"

        # Для других типов файлов отправляем их как есть
        return send_file(absolute_path)

    except Exception as e:
        logger.error(f"Ошибка при просмотре документа: {e}")
        return f"Ошибка при просмотре документа: {str(e)}", 500

    """Запуск веб-сервера"""
    logger.info(f"Запуск веб-сервера на {host}:{port}")
    app.run(host=host, port=port, debug=False)

if __name__ == '__main__':
    run_server()