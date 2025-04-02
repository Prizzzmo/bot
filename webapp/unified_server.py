"""
Объединенный веб-сервер для карты исторических событий и админ-панели
"""

import os
import sys
import json
import logging
import time
import shutil
import threading
import datetime
import re
from flask import Flask, render_template, jsonify, request, send_file, make_response, redirect, url_for

# Настройка логирования
import os
if not os.path.exists('logs'):
    os.makedirs('logs')

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO,
    filename='logs/unified_server.log'
)
logger = logging.getLogger("UnifiedServer")

# Пути к файлам данных
ADMINS_FILE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'admins.json')
BOT_SETTINGS_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'bot_settings.json')
HISTORY_DB_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                             "history_db_generator/russian_history_database.json")

# Подключаем основные компоненты проекта
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.logger import Logger
from src.analytics import AnalyticsService
from src.admin_panel import AdminPanel
from src.config import Config
from src.api_cache import APICache
from src.api_client import APIClient

class UnifiedServer:
    """
    Объединенный сервер для карты и админ-панели
    """

    def __init__(self):
        """Инициализация объединенного сервера"""
        # Создаем логгер
        self.logger = Logger()
        logger.info("Инициализация объединенного сервера")

        # Загружаем конфигурацию
        self.config = Config()

        # Создаем сервисы
        self.analytics_service = AnalyticsService(self.logger)
        self.admin_panel = AdminPanel(self.logger, self.config)

        # Создаем приложение Flask
        template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates')
        static_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'static')

        # Проверяем наличие директорий
        if not os.path.exists(template_path):
            logger.warning(f"Директория шаблонов не найдена: {template_path}")
            os.makedirs(template_path, exist_ok=True)

        if not os.path.exists(static_path):
            logger.warning(f"Директория статических файлов не найдена: {static_path}")
            os.makedirs(static_path, exist_ok=True)

        self.app = Flask(__name__, 
                         template_folder=template_path,
                         static_folder=static_path)


        # Предзагрузка данных
        self.events_data = None
        self.admins_data = None
        self._preload_historical_data()

        # Настраиваем маршруты
        self._setup_routes()

        logger.info("Объединенный сервер инициализирован")

    def _create_demo_events(self):
        """
        Создает демонстрационные исторические события, если основная база данных не может быть загружена

        Returns:
            List: Список демонстрационных исторических событий
        """
        logger.info("Создание демонстрационных исторических событий")
        demo_events = [
            {
                "id": "demo1",
                "title": "Крещение Руси",
                "date": "988 год",
                "description": "Принятие христианства в качестве государственной религии Киевской Руси при князе Владимире.",
                "location": {
                    "lat": 50.4501,
                    "lng": 30.5234,
                    "name": "Киев"
                },
                "category": "Культура и религия",
                "topic": "Христианизация"
            },
            {
                "id": "demo2",
                "title": "Куликовская битва",
                "date": "1380 год",
                "description": "Сражение между объединённым русским войском под командованием московского князя Дмитрия Донского и войском темника Золотой Орды Мамая.",
                "location": {
                    "lat": 53.6764,
                    "lng": 38.6619,
                    "name": "Куликово поле"
                },
                "category": "Войны и сражения",
                "topic": "Освобождение от монголо-татарского ига"
            },
            {
                "id": "demo3",
                "title": "Основание Санкт-Петербурга",
                "date": "1703 год",
                "description": "Основание города Санкт-Петербурга Петром I, который впоследствии стал столицей Российской империи.",
                "location": {
                    "lat": 59.9343,
                    "lng": 30.3351,
                    "name": "Санкт-Петербург"
                },
                "category": "Основание городов",
                "topic": "Петровские реформы"
            },
            {
                "id": "demo4",
                "title": "Отечественная война 1812 года",
                "date": "1812 год",
                "description": "Война между Российской и Французской империями на территории России в 1812 году.",
                "location": {
                    "lat": 55.7522,
                    "lng": 37.6156,
                    "name": "Москва"
                },
                "category": "Войны и сражения",
                "topic": "Наполеоновские войны"
            },
            {
                "id": "demo5",
                "title": "Отмена крепостного права",
                "date": "1861 год",
                "description": "Крестьянская реформа, начатая в 1861 году, упразднила крепостное право в Российской империи.",
                "location": {
                    "lat": 55.7522,
                    "lng": 37.6156,
                    "name": "Москва"
                },
                "category": "Социальные реформы",
                "topic": "Великие реформы"
            },
            {
                "id": "demo6",
                "title": "Октябрьская революция",
                "date": "1917 год",
                "description": "Вооружённое восстание, организованное большевиками, в результате которого было свергнуто Временное правительство и к власти пришло правительство, сформированное II Всероссийским съездом Советов.",
                "location": {
                    "lat": 59.9343,
                    "lng": 30.3351,
                    "name": "Петроград"
                },
                "category": "Революции и перевороты",
                "topic": "Революция 1917 года"
            },
            {
                "id": "demo7",
                "title": "Великая Отечественная война",
                "date": "1941-1945 годы",
                "description": "Война Советского Союза против нацистской Германии и её союзников в рамках Второй мировой войны.",
                "location": {
                    "lat": 55.7522,
                    "lng": 37.6156,
                    "name": "Москва"
                },
                "category": "Войны и сражения",
                "topic": "Вторая мировая война"
            },
            {
                "id": "demo8",
                "title": "Первый полёт человека в космос",
                "date": "1961 год",
                "description": "Первый в мире полёт человека в космическое пространство, совершенный Юрием Гагариным на корабле «Восток-1».",
                "location": {
                    "lat": 51.2754,
                    "lng": 45.9993,
                    "name": "Байконур"
                },
                "category": "Научные достижения",
                "topic": "Космическая программа"
            },
            {
                "id": "demo9",
                "title": "Распад СССР",
                "date": "1991 год",
                "description": "Процессы, которые привели к прекращению существования СССР и появлению независимых государств на постсоветском пространстве.",
                "location": {
                    "lat": 55.7522,
                    "lng": 37.6156,
                    "name": "Москва"
                },
                "category": "Политические реформы",
                "topic": "Перестройка"
            },
            {
                "id": "demo10",
                "title": "Принятие Конституции РФ",
                "date": "1993 год",
                "description": "Принятие всенародным голосованием новой Конституции Российской Федерации.",
                "location": {
                    "lat": 55.7522,
                    "lng": 37.6156,
                    "name": "Москва"
                },
                "category": "Политические реформы",
                "topic": "Современная Россия"
            }
        ]
        return demo_events

    def _preload_historical_data(self):
        """Предварительная загрузка исторических данных"""
        try:
            logger.info(f"Попытка загрузки исторических данных из: {HISTORY_DB_PATH}")
            if os.path.exists(HISTORY_DB_PATH):
                # Проверяем размер файла
                file_size = os.path.getsize(HISTORY_DB_PATH)
                logger.info(f"Файл базы данных найден. Размер: {file_size} байт")

                if file_size == 0:
                    logger.warning("Файл базы данных пуст. Использую пустой набор событий.")
                    self.events_data = {"events": []}
                    return

                try:
                    with open(HISTORY_DB_PATH, 'r', encoding='utf-8') as f:
                        content = f.read().strip()
                        if not content:
                            logger.warning("Содержимое файла пусто. Использую пустой набор событий.")
                            self.events_data = {"events": []}
                            return

                        # Пытаемся загрузить JSON
                        self.events_data = json.loads(content)
                        events_count = len(self.events_data.get('events', []))
                        logger.info(f"Успешно загружено {events_count} исторических событий")
                except json.JSONDecodeError as je:
                    logger.error(f"Ошибка декодирования JSON: {je}.  File content: {content[:100]}...")
                    self.events_data = {"events": []}
                    return #Added return to stop further execution after error
            else:
                logger.warning(f"Файл базы данных не найден: {HISTORY_DB_PATH}")
                # Попробуем найти другие JSON файлы с историческими данными
                for alternative_file in ['historical_events.json', 'history_db_generator/events.json']:
                    alternative_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), alternative_file)
                    if os.path.exists(alternative_path):
                        logger.info(f"Найден альтернативный файл данных: {alternative_path}")
                        try:
                            with open(alternative_path, 'r', encoding='utf-8') as f:
                                self.events_data = json.load(f)
                            logger.info(f"Загружено {len(self.events_data.get('events', []))} событий из альтернативного файла")
                            return
                        except Exception as alt_e:
                            logger.error(f"Ошибка при загрузке альтернативного файла: {alt_e}")
                            return #Added return to stop further execution after error

                # Если не нашли альтернативных файлов, используем пустой набор
                self.events_data = {"events": []}
        except Exception as e:
            logger.error(f"Ошибка при загрузке исторических данных: {str(e)}")
            import traceback
            logger.error(f"Трассировка: {traceback.format_exc()}")
            self.events_data = {"events": []}


    def _clean_event_data(self, events):
        """
        Очищает и форматирует данные событий для API

        Args:
            events: Список событий

        Returns:
            Очищенный список событий
        """
        cleaned_events = []

        for event in events:
            # Проверка типа данных события
            if not isinstance(event, dict):
                logger.warning(f"Пропуск события неверного типа: {type(event)}")
                continue

            # Пропускаем события без заголовка или даты
            if not event.get('title') or not event.get('date'):
                continue

            # Пропускаем события без местоположения или координат
            if 'location' not in event:
                continue

            # Проверка структуры location
            location = event.get('location')
            if not isinstance(location, dict):
                continue

            # Проверка наличия координат
            if not location.get('lat') or not location.get('lng'):
                continue

            # Проверяем и очищаем описание
            description = event.get('description', '')
            if description:
                # Удаляем лишние пробелы и переносы строк
                description = re.sub(r'\s+', ' ', description).strip()
                # Удаляем лишние звездочки, которые не являются частью маркированного списка или жирного текста
                description = re.sub(r'(?<!\*)\*(?!\s|\*)', '', description)

            # Создаем очищенный объект события
            clean_event = {
                'id': event.get('id', ''),
                'title': event.get('title', '').strip(),
                'date': event.get('date', '').strip(),
                'description': description,
                'location': location,
                'category': event.get('category', '').strip(),
                'topic': event.get('topic', '').strip(),
                'century': self._extract_century(event.get('date', ''))
            }

            cleaned_events.append(clean_event)

        return cleaned_events

    def _extract_century(self, date_str):
        """
        Извлекает век из строки даты

        Args:
            date_str: Строка с датой

        Returns:
            Номер века или 0, если не удалось извлечь
        """
        if not date_str:
            return 0

        # Ищем 4-значный год
        year_match = re.search(r'\b(\d{4})\b', date_str)
        if year_match:
            year = int(year_match.group(1))
            return (year // 100) + 1

        # Ищем любое число, которое может быть годом
        year_match = re.search(r'\b(\d+)\b', date_str)
        if year_match:
            year = int(year_match.group(1))
            # Проверяем, что это может быть год (от 800 до текущего года + 100)
            if 800 <= year <= datetime.datetime.now().year + 100:
                return (year // 100) + 1

        return 0

    def _load_admins(self):
        """Загружает список администраторов из файла"""
        try:
            # Если у нас уже есть данные в кэше и они актуальны, используем их
            if self.admins_data is not None:
                return self.admins_data

            if os.path.exists(ADMINS_FILE_PATH):
                with open(ADMINS_FILE_PATH, 'r', encoding='utf-8') as f:
                    self.admins_data = json.load(f)
                return self.admins_data

            # Если файл не существует, создаем структуру по умолчанию
            self.admins_data = {"admin_ids": [], "super_admin_ids": []}
            return self.admins_data
        except Exception as e:
            logger.error(f"Ошибка при загрузке списка администраторов: {e}")
            return {"admin_ids": [], "super_admin_ids": []}

    def _save_admins(self, admins):
        """Сохраняет список администраторов в файл"""
        try:
            # Используем атомарную операцию записи через временный файл
            temp_file = f"{ADMINS_FILE_PATH}.tmp"
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(admins, f, indent=4)

            # Заменяем основной файл только после успешной записи
            os.replace(temp_file, ADMINS_FILE_PATH)

            # Обновляем кэш
            self.admins_data = admins

            return True
        except Exception as e:
            logger.error(f"Ошибка при сохранении списка администраторов: {e}")
            return False

    def _load_bot_settings(self):
        """Загружает настройки бота из файла"""
        try:
            if os.path.exists(BOT_SETTINGS_PATH):
                with open(BOT_SETTINGS_PATH, 'r', encoding='utf-8') as f:
                    return json.load(f)

            # Если файл не существует, создаем настройки по умолчанию
            return {
                "auto_update_topics": True,
                "collect_statistics": True,
                "developer_mode": False,
                "private_mode": False,
                "notification_level": "all",
                "api_model": "gemini-pro",
                "cache_duration": 24,
                "clear_cache_on_startup": True
            }
        except Exception as e:
            logger.error(f"Ошибка при загрузке настроек бота: {e}")
            return {}

    def _save_bot_settings(self, settings):
        """Сохраняет настройки бота в файл"""
        try:
            # Используем атомарную операцию записи через временный файл
            temp_file = f"{BOT_SETTINGS_PATH}.tmp"
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(settings, f, indent=4)

            # Заменяем основной файл только после успешной записи
            os.replace(temp_file, BOT_SETTINGS_PATH)

            return True
        except Exception as e:
            logger.error(f"Ошибка при сохранении настроек бота: {e}")
            return False

    def _get_last_logs(self, lines=100, level='all'):
        """Получает последние строки из файла логов"""
        try:
            log_files = []
            log_dir = "logs"

            # Проверяем, существует ли директория логов
            if os.path.exists(log_dir) and os.path.isdir(log_dir):
                # Получаем список файлов логов
                files = os.listdir(log_dir)
                log_files = [os.path.join(log_dir, f) for f in files if f.startswith("bot_log_")]

            # Если директории нет или нет файлов логов, проверяем в корневой директории
            if not log_files:
                files = [f for f in os.listdir() if f.startswith("bot_log_")]
                log_files = files

            # Сортируем файлы по дате изменения (новые в начале)
            log_files.sort(key=lambda x: os.path.getmtime(x), reverse=True)

            if not log_files:
                return ["Файлы логов не найдены"]

            # Берем самый свежий файл логов
            latest_log = log_files[0]

            with open(latest_log, 'r', encoding='utf-8') as f:
                # Читаем все строки файла
                all_lines = f.readlines()

            # Фильтруем по уровню, если указан не 'all'
            if level != 'all':
                level_upper = level.upper()
                filtered_lines = [line for line in all_lines if f" - {level_upper} - " in line]
            else:
                filtered_lines = all_lines

            # Берем последние N строк
            return filtered_lines[-lines:]
        except Exception as e:
            logger.error(f"Ошибка при чтении файла логов: {e}")
            return [f"Ошибка при чтении логов: {e}"]

    def _setup_routes(self):
        """Настраивает маршруты для объединенного сервера"""

        # ==================== КАРТА ИСТОРИИ ====================

        @self.app.route('/')
        def index():
            """Главная страница с картой истории"""
            # Получаем статистику для приветственного баннера
            events_count = len(self.events_data.get('events', [])) if self.events_data else 0

            # Получаем уникальные категории
            categories = set()
            if self.events_data and 'events' in self.events_data:
                for event in self.events_data['events']:
                    if event.get('category'):
                        categories.add(event.get('category'))

            categories_count = len(categories)

            return render_template('index.html', 
                                  title="История России на карте",
                                  events_count=events_count,
                                  categories_count=categories_count)

        @self.app.route('/api/historical-events')
        def get_historical_events():
            """API для получения исторических данных из базы"""
            try:
                # Проверяем, загружены ли данные
                if self.events_data is None:
                    logger.info("Данные не загружены, выполняем загрузку")
                    self._preload_historical_data()

                # Получаем события из базы данных
                events = self.events_data.get('events', [])
                logger.info(f"Всего событий в базе данных: {len(events)}")

                # Если событий нет, попробуем загрузить демо-данные
                if len(events) == 0:
                    logger.warning("События не найдены, создаем демо-данные")
                    # Создаем тестовые данные для демонстрации
                    events = self._create_demo_events()
                    logger.info(f"Создано {len(events)} демо-событий")

                # Очищаем и форматируем данные
                filtered_events = self._clean_event_data(events)
                logger.info(f"После фильтрации осталось {len(filtered_events)} событий")

                # Фильтрация по параметрам запроса (если они есть)
                category = request.args.get('category')
                century = request.args.get('century')

                if category:
                    filtered_events = [e for e in filtered_events if e.get('category') == category]

                if century:
                    try:
                        century_int = int(century)
                        filtered_events = [e for e in filtered_events if e.get('century') == century_int]
                    except ValueError:
                        pass

                logger.info(f"Отправляется {len(filtered_events)} событий")
                return jsonify(filtered_events)
            except Exception as e:
                logger.error(f"Ошибка при получении исторических данных: {e}")
                import traceback
                logger.error(f"Трассировка: {traceback.format_exc()}")
                return jsonify({'error': str(e)}), 500

        @self.app.route('/api/categories')
        def get_categories():
            """API для получения списка категорий событий"""
            try:
                events = self.events_data.get('events', [])
                categories = sorted(list(set(e.get('category') for e in events if e.get('category'))))

                return jsonify(categories)
            except Exception as e:
                logger.error(f"Ошибка при получении категорий: {e}")
                return jsonify({'error': str(e)}), 500

        @self.app.route('/api/event-details', methods=['POST'])
        def get_event_details():
            """API для получения информации о событии через Gemini API"""
            try:
                # Получаем данные из запроса
                data = request.json

                if not data or not data.get('title'):
                    return jsonify({'error': 'Недостаточно данных о событии'}), 400

                # Получаем API ключ
                api_key = None
                try:
                    from gemini_api_keys import GEMINI_API_KEYS
                    if GEMINI_API_KEYS:
                        api_key = GEMINI_API_KEYS[0]
                except ImportError:
                    logger.error("Не удалось импортировать GEMINI_API_KEYS")

                if not api_key:
                    try:
                        import os
                        api_key = os.environ.get('GEMINI_API_KEY')
                    except:
                        pass

                if not api_key:
                    return jsonify({
                        'content': 'Не удалось получить доступ к API Gemini. Пожалуйста, проверьте настройки API ключей.'
                    }), 200

                # Создаем объекты для работы с API
                api_cache = APICache(self.logger)
                api_client = APIClient(api_key, api_cache, self.logger)

                # Формируем промпт для Gemini
                event_title = data.get('title', '')
                event_date = data.get('date', '')
                event_category = data.get('category', '')
                event_location = data.get('location', '')
                is_brief = data.get('isBrief', True)

                if is_brief:
                    prompt = f"""
                    Предоставь краткую историческую информацию о следующем событии из истории России:

                    Название: {event_title}
                    Дата: {event_date}
                    Категория: {event_category}
                    Место: {event_location}

                    Ответ должен быть кратким (не более 200 слов), но содержательным. 
                    Выдели 3-4 ключевых факта о событии и его значении. 
                    Используй маркированный список для лучшей читаемости.
                    """
                else:
                    prompt = f"""
                    Предоставь подробную историческую информацию о следующем событии из истории России:

                    Название: {event_title}
                    Дата: {event_date}
                    Категория: {event_category}
                    Место: {event_location}

                    Пожалуйста, структурируй ответ следующим образом:
                    1. Исторический контекст (что происходило в России в это время)
                    2. Подробное описание события
                    3. Ключевые участники
                    4. Причины и предпосылки
                    5. Последствия и историческое значение

                    Используй только проверенные исторические факты. Ответ должен быть информативным и подробным.
                    """

                try:
                    # Инициализация API клиента если нужно
                    if not api_client.is_initialized():
                        api_client.initialize()

                    # Запрос к Gemini API
                    response = api_client.ask_grok(prompt, use_cache=True)

                    return jsonify({
                        'content': response
                    })
                except Exception as api_error:
                    logger.error(f"Ошибка при запросе к Gemini API: {api_error}")
                    return jsonify({
                        'content': f"Произошла ошибка при получении информации: {str(api_error)}"
                    }), 200

            except Exception as e:
                logger.error(f"Ошибка при обработке запроса о деталях события: {e}")
                return jsonify({
                    'content': 'Произошла ошибка при обработке запроса. Пожалуйста, попробуйте позже.'
                }), 200

        @self.app.route('/api/generate-report', methods=['POST'])
        def generate_report():
            """API для генерации подробного реферата в формате Word"""
            try:
                # Получаем данные из запроса
                data = request.json

                if not data or not data.get('title'):
                    return jsonify({'error': 'Недостаточно данных о событии'}), 400

                # Импортируем необходимые библиотеки
                import sys
                import os
                from io import BytesIO
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                # Получаем API ключ
                api_key = None
                try:
                    from gemini_api_keys import GEMINI_API_KEYS
                    if GEMINI_API_KEYS:
                        api_key = GEMINI_API_KEYS[0]
                except ImportError:
                    logger.error("Не удалось импортировать GEMINI_API_KEYS")

                if not api_key:
                    try:
                        import os
                        api_key = os.environ.get('GEMINI_API_KEY')
                    except:
                        pass

                if not api_key:
                    return jsonify({'error': 'API ключ не найден'}), 500

                # Создаем объекты для работы с API
                api_cache = APICache(self.logger)
                api_client = APIClient(api_key, api_cache, self.logger)

                # Формируем промпт для Gemini для получения подробной информации
                event_title = data.get('title', '')
                event_date = data.get('date', '')
                event_category = data.get('category', '')
                event_location = data.get('location', '')

                prompt = f"""
                Напиши подробный исторический реферат о следующем событии из истории России:

                Название: {event_title}
                Дата: {event_date}
                Категория: {event_category}
                Место: {event_location}

                Реферат должен содержать следующие разделы:
                1. Введение и исторический контекст (что происходило в России в это время)
                2. Хронология и детальное описание события
                3. Ключевые исторические личности, их роли и биографические данные
                4. Причины и предпосылки события
                5. Непосредственные последствия
                6. Историческое значение и влияние на дальнейший ход истории России
                7. Историографический анализ и различные точки зрения историков
                8. Заключение

                Реферат должен быть максимально подробным, академическим и опираться на исторические источники.
                """

                try:
                    # Инициализация API клиента если нужно
                    if not api_client.is_initialized():
                        api_client.initialize()

                    # Запрос к Gemini API
                    detailed_content = api_client.ask_grok(prompt, use_cache=True)

                    # Создаем документ Word
                    doc = Document()

                    # Устанавливаем стили
                    title_style = doc.styles['Title']
                    title_style.font.size = Pt(16)
                    title_style.font.bold = True

                    heading_style = doc.styles['Heading 1']
                    heading_style.font.size = Pt(14)
                    heading_style.font.bold = True

                    # Заголовок документа
                    doc.add_heading(event_title, 0)

                    # Информация о событии
                    info_paragraph = doc.add_paragraph()
                    info_paragraph.add_run(f"Дата: {event_date}\n").bold = True
                    info_paragraph.add_run(f"Категория: {event_category}\n").bold = True
                    info_paragraph.add_run(f"Место: {event_location}").bold = True

                    # Добавляем разделительную линию
                    doc.add_paragraph("_" * 50)

                    # Парсим и добавляем содержимое
                    content_lines = detailed_content.split('\n')
                    current_heading_level = 0

                    for i, line in enumerate(content_lines):
                        line = line.strip()

                        if not line:
                            doc.add_paragraph()
                            continue

                        # Определяем, является ли строка заголовком
                        heading_level = 0
                        if line.startswith('# '):
                            heading_level = 1
                            heading_text = line[2:].strip()
                        elif line.startswith('## '):
                            heading_level = 2
                            heading_text = line[3:].strip()
                        elif line.startswith('### '):
                            heading_level = 3
                            heading_text = line[4:].strip()
                        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                            heading_level = 2
                            heading_text = line[3:].strip()

                        # Обрабатываем заголовки
                        if heading_level > 0:
                            if heading_level == 1:
                                doc.add_heading(heading_text, level=1)
                            elif heading_level == 2:
                                doc.add_heading(heading_text, level=2)
                            elif heading_level == 3:
                                doc.add_heading(heading_text, level=3)

                            current_heading_level = heading_level
                            continue

                        # Обрабатываем обычный текст
                        if line.startswith('- ') or line.startswith('* '):
                            # Маркированный список                            p = doc.add_paragraph(line[2:], style='List Bullet')
                        else:
                            # Обычный текст
                            p = doc.add_paragraph(line)

                                        # Сохраняем документ в память
                    file_buffer = BytesIO()
                    doc.save(file_buffer)
                    file_buffer.seek(0)

                    # Возвращаем файл
                    return send_file(
                        file_buffer,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f"{event_title.replace('/', '_')}_реферат.docx"
                    )

                except Exception as api_error:
                    logger.error(f"Ошибка при генерации реферата: {api_error}")
                    return jsonify({'error': str(api_error)}), 500

            except Exception as e:
                logger.error(f"Ошибка при обработке запроса на генерацию реферата: {e}")
                return jsonify({'error': str(e)}), 500

            except Exception as e:
                logger.error(f"Ошибка при обработке запроса на генерацию реферата: {e}")
                return jsonify({'error': str(e)}), 500

        # ==================== АДМИН-ПАНЕЛЬ ====================

        @self.app.route('/login')
        def login_page():
            """Страница авторизации"""
            return render_template('login.html', title="Вход в админ-панель")

        @self.app.route('/admin-panel')
        def admin_panel():
            """Главная страница админ-панели"""
            # Проверка авторизации через cookie
            user_id = request.cookies.get('admin_id')
            if not user_id:
                # Если пользователь не авторизован, перенаправляем на страницу входа
                return redirect('/login')

            return render_template('admin_panel.html', title="Панель администратора")

        # -------------------- Аутентификация --------------------

        @self.app.route('/api/admin/check-auth', methods=['GET'])
        def api_admin_check_auth():
            """Проверка аутентификации администратора"""
            user_id = request.cookies.get('admin_id')

            if not user_id:
                return jsonify({"authenticated": False})

            try:
                user_id = int(user_id)
                admins = self._load_admins()

                # Проверяем, является ли пользователь администратором
                is_admin = user_id in admins.get("admin_ids", []) or user_id in admins.get("super_admin_ids", [])
                is_super_admin = user_id in admins.get("super_admin_ids", [])

                if is_admin:
                    return jsonify({
                        "authenticated": True,
                        "user": {
                            "id": user_id,
                            "is_super_admin": is_super_admin
                        }
                    })

                return jsonify({"authenticated": False})
            except Exception as e:
                logger.error(f"Ошибка при проверке аутентификации: {e}")
                return jsonify({"authenticated": False})

        @self.app.route('/api/admin/login', methods=['POST'])
        def api_admin_login():
            """Вход администратора"""
            try:
                data = request.get_json()
                logger.info(f"Получен запрос на авторизацию: {data}")

                # Проверяем тип входа: по ID или паролю
                if 'admin_id' in data:
                    admin_id = int(data.get('admin_id', 0))
                    logger.info(f"Попытка входа по ID: {admin_id}")

                    if not admin_id:
                        logger.warning("ID администратора не указан")
                        return jsonify({"success": False, "message": "ID администратора не указан"})

                    admins = self._load_admins()
                    logger.info(f"Загруженные данные администраторов: {admins}")

                    # Проверяем, является ли пользователь администратором
                    is_admin = admin_id in admins.get("admin_ids", []) or admin_id in admins.get("super_admin_ids", [])
                    is_super_admin = admin_id in admins.get("super_admin_ids", [])

                    # Проверка является ли пользователь администратором
                    if not is_admin:
                        logger.warning(f"Неверный ID администратора: {admin_id}")
                        return jsonify({"success": False, "message": "Неверный ID администратора"})

                    if is_admin:
                        logger.info(f"Пользователь {admin_id} авторизован как {'супер-' if is_super_admin else ''}администратор")
                        response = jsonify({
                            "success": True,
                            "user": {
                                "id": admin_id,
                                "is_super_admin": is_super_admin
                            }
                        })
                        response.set_cookie('admin_id', str(admin_id), max_age=86400)  # 24 часа
                        logger.info(f"Успешный вход администратора по ID: {admin_id}")
                        return response

                    logger.warning(f"Неверный ID администратора: {admin_id}")
                    return jsonify({"success": False, "message": "Неверный ID администратора"})

                elif 'admin_password' in data:
                    admin_password = data.get('admin_password', '')
                    logger.info(f"Попытка входа по паролю")

                    if not admin_password:
                        logger.warning("Пароль администратора не указан")
                        return jsonify({"success": False, "message": "Пароль администратора не указан"})

                    # Проверяем пароль (в реальном проекте должно быть безопасное хранение)
                    correct_password = "nnkhqjm"  # Фиксированный пароль для админ-панели

                    if admin_password == correct_password:
                        # При успешной авторизации используем ID супер-администратора
                        admins = self._load_admins()
                        admin_id = admins.get("super_admin_ids", [7225056628])[0]

                        logger.info(f"Успешный вход администратора по паролю с ID: {admin_id}")

                        response = jsonify({
                            "success": True,
                            "user": {
                                "id": admin_id,
                                "is_super_admin": True
                            }
                        })
                        response.set_cookie('admin_id', str(admin_id), max_age=86400)  # 24 часа
                        return response

                    logger.warning(f"Неудачная попытка входа администратора с паролем")
                    return jsonify({"success": False, "message": "Неверный пароль администратора"})

                logger.warning(f"Неверный запрос авторизации: {data}")
                return jsonify({"success": False, "message": "Неверный запрос авторизации"})
            except Exception as e:
                logger.error(f"Ошибка при авторизации администратора: {e}", exc_info=True)
                return jsonify({"success": False, "message": f"Ошибка при авторизации: {str(e)}"})

        @self.app.route('/api/admin/logout', methods=['POST'])
        def api_admin_logout():
            """Выход администратора"""
            response = jsonify({"success": True})
            response.delete_cookie('admin_id')
            return response

        # -------------------- Управление администраторами --------------------

        @self.app.route('/api/admin/admins', methods=['GET'])
        def api_admin_get_admins():
            """Получение списка администраторов"""
            try:
                # Проверка аутентификации
                user_id = request.cookies.get('admin_id')
                if not user_id:
                    return jsonify({"error": "Требуется авторизация"}), 403

                admins = self._load_admins()

                # Форматируем данные для фронтенда
                admin_list = []

                # Добавляем супер-админов
                for admin_id in admins.get("super_admin_ids", []):
                    admin_list.append({
                        "id": admin_id,
                        "is_super": True,
                        "type": "Супер-администратор"
                    })

                # Добавляем обычных админов
                for admin_id in admins.get("admin_ids", []):
                    admin_list.append({
                        "id": admin_id,
                        "is_super": False,
                        "type": "Администратор"
                    })

                return jsonify({"admins": admin_list})
            except Exception as e:
                logger.error(f"Ошибка при получении списка администраторов: {e}")
                return jsonify({"error": str(e)}), 500

        @self.app.route('/api/admin/add-admin', methods=['POST'])
        def api_admin_add_admin():
            """Добавление нового администратора"""
            try:
                # Проверка аутентификации
                user_id = request.cookies.get('admin_id')
                if not user_id:
                    return jsonify({"success": False, "message": "Требуется авторизация"}), 403

                user_id = int(user_id)
                admins = self._load_admins()

                # Проверяем, является ли пользователь супер-администратором
                if user_id not in admins.get("super_admin_ids", []):
                    return jsonify({"success": False, "message": "Требуются права супер-администратора"}), 403

                data = request.get_json()
                admin_id = int(data.get('admin_id', 0))
                is_super = data.get('is_super', False)

                if not admin_id:
                    return jsonify({"success": False, "message": "ID администратора не указан"})

                # Проверяем, не существует ли уже такой админ
                if admin_id in admins.get("admin_ids", []) or admin_id in admins.get("super_admin_ids", []):
                    return jsonify({"success": False, "message": "Пользователь уже является администратором"})

                # Добавляем нового администратора
                if is_super:
                    admins.setdefault("super_admin_ids", []).append(admin_id)
                    logger.info(f"Добавлен супер-админ: {admin_id}, добавил: {user_id}")
                else:
                    admins.setdefault("admin_ids", []).append(admin_id)
                    logger.info(f"Добавлен админ: {admin_id}, добавил: {user_id}")

                # Сохраняем изменения
                self._save_admins(admins)

                return jsonify({"success": True})
            except Exception as e:
                logger.error(f"Ошибка при добавлении администратора: {e}")
                return jsonify({"success": False, "message": f"Ошибка при добавлении администратора: {str(e)}"})

        # -------------------- Остальные маршруты админ-панели --------------------

        # API для получения настроек бота
        @self.app.route('/api/admin/settings', methods=['GET'])
        def api_admin_get_settings():
            """Получение настроек бота"""
            try:
                # Проверка аутентификации
                user_id = request.cookies.get('admin_id')
                if not user_id:
                    return jsonify({"error": "Требуется авторизация"}), 403

                # Используем настройки из админ-панели
                if self.admin_panel:
                    try:
                        settings = self.admin_panel._get_bot_settings()
                        return jsonify(settings)
                    except Exception as e:
                        logger.warning(f"Не удалось получить настройки из админ-панели: {e}")

                # Запасной вариант: загрузка из файла
                settings = self._load_bot_settings()
                return jsonify(settings)
            except Exception as e:
                logger.error(f"Ошибка при получении настроек бота: {e}")
                return jsonify({"error": str(e)}), 500

        # API для сохранения настроек бота
        @self.app.route('/api/admin/save-settings', methods=['POST'])
        def api_admin_save_settings():
            """Сохранение настроек бота"""
            try:
                # Проверка аутентификации
                user_id = request.cookies.get('admin_id')
                if not user_id:
                    return jsonify({"success": False, "message": "Требуется авторизация"}), 403

                user_id = int(user_id)
                admins = self._load_admins()

                # Проверяем, является ли пользователь супер-администратором
                if user_id not in admins.get("super_admin_ids", []):
                    return jsonify({"success": False, "message": "Требуются права супер-администратора"}), 403

                settings = request.get_json()

                # Сохраняем настройки в файл
                self._save_bot_settings(settings)

                logger.info(f"Настройки бота обновлены администратором: {user_id}")
                return jsonify({"success": True})
            except Exception as e:
                logger.error(f"Ошибка при сохранении настроек бота: {e}")
                return jsonify({"success": False, "message": f"Ошибка при сохранении настроек: {str(e)}"})

        # API для получения статистики
        @self.app.route('/api/admin/stats', methods=['GET'])
        def api_admin_stats():
            """Получение статистики бота"""
            try:
                # Проверка аутентификации
                user_id = request.cookies.get('admin_id')
                if not user_id:
                    return jsonify({"error": "Требуется авторизация"}), 403

                # Получаем статистику из сервиса аналитики
                if self.analytics_service:
                    try:
                        stats = self.analytics_service.get_overall_stats()
                        return jsonify(stats)
                    except Exception as e:
                        logger.warning(f"Ошибка при получении статистики: {e}")

                # Заглушка для статистики
                return jsonify({
                    "user_count": 42,
                    "message_count": 1337,
                    "uptime": "3 дня 7 часов",
                    "bot_starts": 25,
                    "topic_requests": 73,
                    "completed_tests": 18
                })
            except Exception as e:
                logger.error(f"Ошибка при получении статистики: {e}")
                return jsonify({"error": str(e)}), 500

        # API для получения логов
        @self.app.route('/api/admin/logs', methods=['GET'])
        def api_admin_logs():
            """Получение логов"""
            try:
                # Проверка аутентификации
                user_id = request.cookies.get('admin_id')
                if not user_id:
                    return jsonify({"error": "Требуется авторизация"}), 403

                # Получаем параметры
                level = request.args.get('level', 'all')
                limit = int(request.args.get('limit', 100))

                # Получаем логи
                logs = self._get_last_logs(limit, level)

                return jsonify(logs)
            except Exception as e:
                logger.error(f"Ошибка при получении логов: {e}")
                return jsonify({"error": str(e)}), 500

        # Здесь добавьте остальные маршруты из оригинального admin_server.py
        # ...

def run_unified_server(host='0.0.0.0', port=None):
    """
    Запускает объединенный сервер

    Args:
        host: Хост для запуска сервера
        port: Порт для запуска сервера
    """
    try:
        # Получаем порт из переменных окружения Replit или используем значение по умолчанию
        if port is None:
            port = int(os.environ.get('PORT', 8080))

        # Создаем и запускаем сервер
        server = UnifiedServer()
        replit_url = os.environ.get('REPL_SLUG', None)
        replit_owner = os.environ.get('REPL_OWNER', None)

        if replit_url and replit_owner:
            public_url = f"https://{replit_url}.{replit_owner}.repl.co"
        else:
            public_url = None

        logger.info(f"Объединенный веб-сервер запущен на порту {port} с URL: {public_url}")
        logger.info(f"Карта истории доступна по адресу: /")
        logger.info(f"Админ-панель доступна по адресу: /admin-panel")

        server.app.run(host=host, port=port, debug=False)
    except Exception as e:
        logger.error(f"Ошибка при запуске объединенного сервера: {e}")
        raise e

if __name__ == "__main__":
    # При запуске модуля напрямую запускаем сервер
    run_unified_server()