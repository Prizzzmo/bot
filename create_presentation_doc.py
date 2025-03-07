
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from io import BytesIO
from PIL import Image

def create_presentation_docx(md_file_path, output_docx_path):
    """
    Преобразует Markdown-файл в документ Word с добавлением изображений.
    
    Args:
        md_file_path (str): Путь к файлу Markdown
        output_docx_path (str): Путь для сохранения документа Word
    """
    # Создаем папку для хранения изображений, если она не существует
    images_dir = 'presentation_images'
    os.makedirs(images_dir, exist_ok=True)
    
    # Создаем новый документ Word
    doc = Document()
    
    # Настраиваем стили документа
    styles = doc.styles
    
    # Настраиваем стиль заголовка
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)
    
    # Настраиваем стиль заголовков разделов
    heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(18)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 102, 204)
    
    # Настраиваем стиль подзаголовков
    heading2_style = styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Calibri'
    heading2_font.size = Pt(16)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 153, 153)
    
    # Читаем содержимое Markdown-файла
    with open(md_file_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # Разбиваем контент на строки
    lines = md_content.split('\n')
    
    # Текущий уровень заголовка
    current_heading_level = 0
    
    # Обрабатываем каждую строку
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Пропускаем пустые строки
        if not line:
            i += 1
            continue
        
        # Обрабатываем заголовки
        if line.startswith('#'):
            heading_level = len(line.split(' ')[0])
            heading_text = line.strip('#').strip()
            
            if heading_level == 1:
                # Основной заголовок (Заголовок документа)
                title = doc.add_paragraph(heading_text, style='Title')
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Добавляем изображение под главным заголовком
                if heading_text == "Образовательный бот по истории России":
                    try:
                        # Изображение для основного заголовка
                        image_url = "https://histrf.ru/uploads/media/default/0001/02/0e4f4e9d11f6bc76c7c69de95b3ab6f3c6fb7d3b.jpeg"
                        response = requests.get(image_url)
                        image_path = os.path.join(images_dir, "russia_history.jpg")
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(response.content)
                        
                        doc.add_picture(image_path, width=Inches(6))
                        picture_paragraph = doc.paragraphs[-1]
                        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"Ошибка при добавлении изображения: {e}")
            
            elif heading_level == 2:
                # Заголовок раздела
                doc.add_heading(heading_text, level=1)
                
                # Добавляем тематические изображения для разделов
                if "Архитектура приложения" in heading_text:
                    try:
                        image_url = "https://miro.medium.com/v2/resize:fit:1200/0*9pGRpgfmnqPEBuFj.png"
                        response = requests.get(image_url)
                        image_path = os.path.join(images_dir, "architecture.jpg")
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(response.content)
                        
                        doc.add_picture(image_path, width=Inches(5))
                        picture_paragraph = doc.paragraphs[-1]
                        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"Ошибка при добавлении изображения: {e}")
                
                elif "Взаимодействие с API Gemini" in heading_text:
                    try:
                        image_url = "https://storage.googleapis.com/gweb-uniblog-publish-prod/images/Gemini_Features_2.max-1000x1000.png"
                        response = requests.get(image_url)
                        image_path = os.path.join(images_dir, "gemini_api.jpg")
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(response.content)
                        
                        doc.add_picture(image_path, width=Inches(5))
                        picture_paragraph = doc.paragraphs[-1]
                        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"Ошибка при добавлении изображения: {e}")
                
                elif "Компоненты системы" in heading_text:
                    try:
                        image_url = "https://cdn-icons-png.flaticon.com/512/2530/2530302.png"
                        response = requests.get(image_url)
                        image_path = os.path.join(images_dir, "components.jpg")
                        
                        with open(image_path, 'wb') as img_file:
                            img_file.write(response.content)
                        
                        doc.add_picture(image_path, width=Inches(4))
                        picture_paragraph = doc.paragraphs[-1]
                        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"Ошибка при добавлении изображения: {e}")
            
            elif heading_level == 3:
                # Подзаголовок
                doc.add_heading(heading_text, level=2)
            
            current_heading_level = heading_level
        
        # Обрабатываем блоки кода
        elif line.startswith('```'):
            # Начало блока кода
            if i + 1 < len(lines):
                code_block = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_block.append(lines[i])
                    i += 1
                
                # Добавляем блок кода с форматированием
                code_text = '\n'.join(code_block)
                code_paragraph = doc.add_paragraph()
                code_run = code_paragraph.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(10)
                code_paragraph.style = doc.styles.add_style('CodeBlock', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                code_paragraph.style.font.name = 'Courier New'
                code_paragraph.style.font.size = Pt(10)
                code_paragraph.paragraph_format.space_before = Pt(6)
                code_paragraph.paragraph_format.space_after = Pt(6)
                code_paragraph.paragraph_format.left_indent = Inches(0.5)
                
                # Создаем серый фон для блока кода
                for paragraph in doc.paragraphs:
                    if paragraph.text == code_text:
                        paragraph.style.paragraph_format.left_indent = Inches(0.5)
                        paragraph.style.paragraph_format.right_indent = Inches(0.5)
                        # Шейдинг настраивается через xml, что выходит за рамки простого примера
        
        # Обрабатываем обычный текст
        else:
            # Обработка маркированных списков
            if line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            # Обработка нумерованных списков
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                num, text = line.split('. ', 1)
                p = doc.add_paragraph(text, style='List Number')
            # Обычный текст
            else:
                # Обработка жирного и курсивного текста
                p = doc.add_paragraph()
                remaining_text = line
                
                # Обработка жирного текста
                while '**' in remaining_text:
                    parts = remaining_text.split('**', 2)
                    if len(parts) >= 3:
                        p.add_run(parts[0])
                        bold_run = p.add_run(parts[1])
                        bold_run.bold = True
                        remaining_text = parts[2] if len(parts) > 2 else ""
                    else:
                        p.add_run(remaining_text)
                        remaining_text = ""
                
                if remaining_text:
                    p.add_run(remaining_text)
        
        i += 1
    
    # Добавляем финальное изображение
    try:
        image_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f3/Flag_of_Russia.svg/800px-Flag_of_Russia.svg.png"
        response = requests.get(image_url)
        image_path = os.path.join(images_dir, "russia_flag.jpg")
        
        with open(image_path, 'wb') as img_file:
            img_file.write(response.content)
        
        doc.add_paragraph()
        doc.add_picture(image_path, width=Inches(5))
        picture_paragraph = doc.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Ошибка при добавлении изображения: {e}")
    
    # Добавляем авторскую информацию в нижний колонтитул
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "© 2025 Образовательный бот по истории России"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Сохраняем документ
    doc.save(output_docx_path)
    print(f"Документ Word успешно создан и сохранен: {output_docx_path}")
    return output_docx_path

if __name__ == "__main__":
    # Создаем презентацию в формате Word
    create_presentation_docx('detailed_presentation.md', 'История_России_подробная_презентация.docx')
